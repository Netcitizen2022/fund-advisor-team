#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基金推荐研究报告生成脚本 v1.0
用途：生成标准化基金推荐Word报告（完整版 + 一页纸摘要版）
依赖：python-docx
使用：python3 generate_report.py --client_name 「张先生」 --risk_level R3 \
          --market_status 积极 --funds_json fund_data.json --output_path ../output/FA-20260625-PI001/
铁律：
  - 全部用Python生成Word，不用JS
  - 中文字符串内引号一律用全角「」『』
  - 文件路径用os.path.join，不硬编码斜杠
"""

import argparse
import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print(「[错误] 缺少 python-docx，请运行：python3 -m pip install python-docx --break-system-packages」)
    sys.exit(1)


# ─────────────────────────────────────────────────────────────
# 颜色与样式常量
# ─────────────────────────────────────────────────────────────
COLOR_PRIMARY = RGBColor(0x1A, 0x37, 0x6C)    # 深蓝色（主标题）
COLOR_ACCENT  = RGBColor(0xE8, 0x5D, 0x04)    # 橙色（强调）
COLOR_BODY    = RGBColor(0x2D, 0x2D, 0x2D)    # 深灰（正文）
COLOR_LIGHT   = RGBColor(0xF5, 0xF5, 0xF5)    # 浅灰（表头背景）

RISK_LABELS = {
    「R1」: 「保守型」,
    「R2」: 「稳健型」,
    「R3」: 「平衡型」,
    「R4」: 「进取型」,
    「R5」: 「激进型」,
}

MARKET_STATUS_DESC = {
    「积极」: 「当前流动性宽松，经济复苏信号明确，政策持续支持，市场整体处于积极状态。」,
    「中性」: 「当前市场信号混杂，流动性环境中性，建议按标准框架配置，保持灵活应对。」,
    「谨慎」: 「当前存在流动性收紧或经济下行压力，建议降低权益仓位，增配防御性资产。」,
}


def set_cell_background(cell, color_hex):
    「「「设置表格单元格背景色」「「」
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement(「w:shd」)
    shd.set(qn(「w:val」), 「clear」)
    shd.set(qn(「w:color」), 「auto」)
    shd.set(qn(「w:fill」), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    「「「添加标题段落」「「」
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = color or COLOR_PRIMARY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or COLOR_PRIMARY
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or COLOR_BODY
    return p


def add_body(doc, text, bold=False, indent=False):
    「「「添加正文段落」「「」
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.bold = bold
    run.font.color.rgb = COLOR_BODY
    return p


def add_divider(doc):
    「「「添加分隔线」「「」
    p = doc.add_paragraph(「─」 * 50)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    return p


def build_fund_table(doc, funds):
    「「「构建基金推荐明细表格」「「」
    headers = [「基金名称」, 「代码」, 「类型」, 「近3年年化」, 「最大回撤」, 「推荐占比」, 「核心亮点」]
    table = doc.add_table(rows=1 + len(funds), cols=len(headers))
    table.style = 「Table Grid」
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_background(cell, 「1A376C」)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

    # 数据行
    for row_idx, fund in enumerate(funds):
        row = table.rows[row_idx + 1]
        values = [
            fund.get(「name」, 「」),
            fund.get(「code」, 「」),
            fund.get(「type」, 「」),
            fund.get(「annual_return_3y」, 「」),
            fund.get(「max_drawdown」, 「」),
            fund.get(「weight」, 「」),
            fund.get(「highlight」, 「」),
        ]
        for col_idx, val in enumerate(values):
            cell = row.cells[col_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = COLOR_BODY
            if row_idx % 2 == 0:
                set_cell_background(cell, 「F8F9FA」)

    return table


def generate_full_report(client_name, risk_level, market_status, funds, output_path):
    「「「生成完整研究报告」「「」
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    report_date = datetime.now().strftime(「%Y年%m月%d日」)
    risk_label  = RISK_LABELS.get(risk_level, risk_level)

    doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(「基金投资组合推荐报告」)
    run.font.size  = Pt(24)
    run.font.bold  = True
    run.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_sub.add_run(f「专属定制 · {client_name} · {report_date}」)
    run2.font.size  = Pt(12)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    add_heading(doc, 「第一章  市场在哪里——我们看到了什么」, level=1)
    add_body(doc, MARKET_STATUS_DESC.get(market_status, 「」))
    doc.add_paragraph()
    add_heading(doc, f「市场研判结论：{market_status}」, level=2, color=COLOR_ACCENT)
    doc.add_paragraph()

    add_heading(doc, 「第二章  什么都不做的代价」, level=1)
    add_body(doc, 「当前一年期存款基准利率持续下行，货币基金收益率走低。」)
    add_body(doc, 「「不做决策」本身就是一个代价高昂的决策。」, bold=True)
    doc.add_paragraph()

    add_heading(doc, 「第三章  我们怎么选——筛选标准透明化」, level=1)
    add_body(doc, f「基于您的{risk_label}风险偏好，从五个维度评分筛选。」)
    doc.add_paragraph()
    add_heading(doc, 「推荐基金明细」, level=2)
    build_fund_table(doc, funds)

    add_heading(doc, 「第四章  组合是怎么搭的」, level=1)
    add_body(doc, f「基于{risk_label}投资者标准框架，采用核心+卫星双层结构。」)
    doc.add_paragraph()

    add_heading(doc, 「第五章  风险在哪里——我们不回避的话题」, level=1)
    add_body(doc, 「我们主动告诉您这个组合可能亏钱的情况：」)
    doc.add_paragraph()

    add_heading(doc, 「第六章  接下来怎么做——明确的行动清单」, level=1)
    add_body(doc, 「建议分批建仓，降低时机选择风险。」)
    add_body(doc, 「下次组合检视时间：3个月后（或发生触发条件时）」, bold=True)
    doc.add_paragraph()

    add_divider(doc)
    p_d = doc.add_paragraph()
    r_d = p_d.add_run(「风险提示：本报告基于历史数据分析，不构成投资建议。」)
    r_d.font.size  = Pt(8)
    r_d.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    os.makedirs(output_path, exist_ok=True)
    date_str  = datetime.now().strftime(「%Y%m%d」)
    filename  = f「基金推荐报告_{client_name}_{date_str}.docx」
    full_path = os.path.join(output_path, filename)
    doc.save(full_path)
    print(f「[OK] 完整报告已生成：{full_path}」)
    return full_path


def generate_summary_sheet(client_name, risk_level, market_status, funds, output_path):
    「「「生成一页纸行动摘要」「「」
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    report_date = datetime.now().strftime(「%Y年%m月%d日」)
    risk_label  = RISK_LABELS.get(risk_level, risk_level)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f「基金组合行动摘要 · {client_name} · {report_date}」)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = COLOR_PRIMARY

    add_divider(doc)
    add_heading(doc, 「推荐基金一览」, level=2)
    build_fund_table(doc, funds)

    os.makedirs(output_path, exist_ok=True)
    date_str  = datetime.now().strftime(「%Y%m%d」)
    filename  = f「基金组合一页纸_{client_name}_{date_str}.docx」
    full_path = os.path.join(output_path, filename)
    doc.save(full_path)
    print(f「[OK] 一页纸摘要已生成：{full_path}」)
    return full_path


def main():
    parser = argparse.ArgumentParser(description=「基金推荐报告生成器 v1.0（已归档）」)
    parser.add_argument(「--client_name」,   required=True)
    parser.add_argument(「--risk_level」,    required=True)
    parser.add_argument(「--market_status」, required=True)
    parser.add_argument(「--funds_json」,    required=True)
    parser.add_argument(「--output_path」,   required=True)
    args = parser.parse_args()

    with open(args.funds_json, 「r」, encoding=「utf-8」) as f:
        data = json.load(f)
    funds = data if isinstance(data, list) else data.get(「funds」, [])

    generate_full_report(args.client_name, args.risk_level, args.market_status, funds, args.output_path)
    generate_summary_sheet(args.client_name, args.risk_level, args.market_status, funds, args.output_path)
    print(「[完成] v1.0 两份文档均已生成。」)


if __name__ == 「__main__」:
    main()
